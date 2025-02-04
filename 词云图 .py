from docx import Document
from docx.enum.text import WD_COLOR_INDEX


def extract_yellow_highlights(input_path, output_path):
    """
    从输入文档中提取黄色高亮文本到新文档
    :param input_path: 输入文档路径
    :param output_path: 输出文档路径
    """
    doc = Document(input_path)
    new_doc = Document()

    for para in doc.paragraphs:
        new_para = None  # 新段落容器

        # 遍历段落中的所有文字块
        for run in para.runs:
            if run.highlight_color == WD_COLOR_INDEX.YELLOW:
                # 只有发现高亮时才创建新段落
                if new_para is None:
                    new_para = new_doc.add_paragraph()

                # 复制文字块并保留格式
                new_run = new_para.add_run(run.text)

                # 复制基础格式
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline

                # 复制字体格式
                new_run.font.name = run.font.name
                if run.font.size:
                    new_run.font.size = run.font.size

                # 保留高亮颜色
                new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

                # 复制文字颜色（可选）
                if run.font.color.rgb:
                    new_run.font.color.rgb = run.font.color.rgb

    new_doc.save(output_path)


# 使用示例
extract_yellow_highlights("input.docx", "output.docx")